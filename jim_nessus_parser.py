#!/usr/bin/env python3
"""
Nessus Parser v1.8
Author: Sir Jimbet (https://github.com/jimbet/)
Export Nessus result to .nessus and parsed it
"""

import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict
import sys
import os
import platform
import subprocess

# CVSS v4 Configuration for severity mapping
CVSS_V4_CONFIG = {
    'critical': {'min': 9.0, 'max': 10.0},
    'high': {'min': 7.0, 'max': 8.9},
    'medium': {'min': 4.0, 'max': 6.9},
    'low': {'min': 0.1, 'max': 3.9},
    'info': {'min': 0.0, 'max': 0.0}
}

# Font Configuration
FONT_CONFIG = {
    'name': 'Arial',  # Change to your preferred font
    'size': 11        # Change to your preferred size
}

def get_severity_from_cvss(cvss_score):
    """Determine severity based on CVSS score using v4 thresholds"""
    try:
        score = float(cvss_score)
        if score >= CVSS_V4_CONFIG['critical']['min']:
            return 'Critical'
        elif score >= CVSS_V4_CONFIG['high']['min']:
            return 'High'
        elif score >= CVSS_V4_CONFIG['medium']['min']:
            return 'Medium'
        elif score > CVSS_V4_CONFIG['low']['min']:
            return 'Low'
        else:
            return 'Info'
    except (ValueError, TypeError):
        return 'Info'

def parse_nessus(file_path):
    """Parse Nessus XML file and extract vulnerabilities"""
    tree = ET.parse(file_path)
    root = tree.getroot()
    
    # Dictionary to store vulnerabilities by severity
    vulns_by_severity = {
        'Critical': [],
        'High': [],
        'Medium': [],
        'Low': [],
        'Info': []
    }
    
    # Parse each ReportHost
    for host in root.findall('.//ReportHost'):
        host_name = host.get('name')
        
        # Parse each ReportItem (vulnerability)
        for item in host.findall('ReportItem'):
            severity = int(item.get('severity', 0))
            
            # Map severity number to text
            severity_map = {
                4: 'Critical',
                3: 'High',
                2: 'Medium',
                1: 'Low',
                0: 'Info'
            }
            severity_text = severity_map.get(severity, 'Info')
            
            # Get CVSS score and use it to determine severity if available
            cvss_score = item.findtext('cvss_base_score', 'N/A')
            if cvss_score != 'N/A':
                severity_text = get_severity_from_cvss(cvss_score)
            
            # Extract vulnerability details
            vuln = {
                'host': host_name,
                'port': item.get('port'),
                'protocol': item.get('protocol'),
                'plugin_id': item.get('pluginID'),
                'plugin_name': item.get('pluginName'),
                'severity': severity_text,
                'description': item.findtext('description', 'N/A'),
                'solution': item.findtext('solution', 'N/A'),
                'synopsis': item.findtext('synopsis', 'N/A'),
                'plugin_output': item.findtext('plugin_output', 'N/A'),
                'cvss_base_score': cvss_score,
                'cvss_vector': item.findtext('cvss_vector', 'N/A'),
                'cvss3_vector': item.findtext('cvss3_vector', 'N/A'),
            }
            
            vulns_by_severity[severity_text].append(vuln)
    
    return vulns_by_severity

def set_cell_background(cell, color_rgb):
    """Set background color for a table cell"""
    shading_elm = OxmlElement('w:shd')
    color_code = '%02x%02x%02x' % color_rgb
    shading_elm.set(qn('w:fill'), color_code)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_font(run, font_name=None, font_size=None, bold=False, italic=False, color=None):
    """Set font properties for a run"""
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.font.bold = bold
    if italic:
        run.font.italic = italic
    if color:
        run.font.color.rgb = color

def set_cell_font(cell, font_name=None, font_size=None, bold=False, italic=False):
    """Set font for all paragraphs in a cell"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_font(run, font_name, font_size, bold, italic)

def create_docx(vulns_by_severity, output_file):
    """Create DOCX report from vulnerabilities"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sir Jimbet - Nessus Vulnerability Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Severity colors - Red, Orange, Yellow, Green, Blue
    severity_colors = {
        'Critical': RGBColor(255, 0, 0),      # Red
        'High': RGBColor(255, 165, 0),        # Orange
        'Medium': RGBColor(255, 255, 0),      # Yellow
        'Low': RGBColor(0, 255, 0),           # Green
        'Info': RGBColor(0, 0, 255)           # Blue
    }
    
    # Summary section
    doc.add_heading('Executive Summary', 1)
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    summary_table.rows[0].cells[0].text = 'Severity'
    summary_table.rows[0].cells[1].text = 'Count'
    
    for idx, severity in enumerate(['Critical', 'High', 'Medium', 'Low', 'Info'], 1):
        summary_table.rows[idx].cells[0].text = severity
        summary_table.rows[idx].cells[1].text = str(len(vulns_by_severity[severity]))
        # Color the severity cell
        run = summary_table.rows[idx].cells[0].paragraphs[0].runs[0]
        run.font.color.rgb = severity_colors[severity]
        run.font.bold = True
    
    doc.add_paragraph()
    
    # Detailed findings by severity - CRITICAL, HIGH, MEDIUM, LOW, INFO order
    for severity in ['Critical', 'High', 'Medium', 'Low', 'Info']:
        vulns = vulns_by_severity[severity]
        
        if not vulns:
            continue
        
        # Add severity heading
        heading = doc.add_heading(f'{severity} Severity Findings ({len(vulns)})', 1)
        heading.runs[0].font.color.rgb = severity_colors[severity]
        
        # Group vulnerabilities by plugin name
        grouped_vulns = defaultdict(list)
        for vuln in vulns:
            grouped_vulns[vuln['plugin_name']].append(vuln)
        
        # Sort grouped vulnerabilities by CVSS score (descending)
        sorted_vulns = sorted(
            grouped_vulns.items(),
            key=lambda x: float(x[1][0]['cvss_base_score']) if x[1][0]['cvss_base_score'] != 'N/A' else 0,
            reverse=True
        )
        
        # Add each unique vulnerability
        for plugin_name, vuln_list in sorted_vulns:
            # Use first occurrence for details
            vuln = vuln_list[0]
            
            # Create vulnerability table with header
            vuln_table = doc.add_table(rows=9, cols=2)
            vuln_table.style = 'Table Grid'
            
            # Set column widths - 20% for first column, 80% for second column
            # Total table width approximately 7.5 inches (standard page width minus margins)
            total_width = Inches(7.5)
            vuln_table.columns[0].width = Inches(1.5)  # 20% of 7.5 inches
            vuln_table.columns[1].width = Inches(6.0)  # 80% of 7.5 inches
            
            # Force column width by setting preferred width on each cell
            for row in vuln_table.rows:
                row.cells[0].width = Inches(1.5)
                row.cells[1].width = Inches(6.0)
            
            # Header row with plugin name (merge cells and color background)
            header_cell = vuln_table.rows[0].cells[0]
            header_cell.merge(vuln_table.rows[0].cells[1])
            header_cell.text = plugin_name
            header_paragraph = header_cell.paragraphs[0]
            header_run = header_paragraph.runs[0]
            header_run.font.bold = True
            header_run.font.size = Pt(14)
            header_run.font.name = FONT_CONFIG['name']
            set_cell_background(header_cell, severity_colors[severity])
            
            # Set font color based on severity (black for yellow background, white for others)
            if severity == 'Medium' or severity == 'Low':
                header_run.font.color.rgb = RGBColor(0, 0, 0)  # Black for yellow background
            else:
                header_run.font.color.rgb = RGBColor(255, 255, 255)  # White for other backgrounds
            
            # Affected Host(s)
            vuln_table.rows[1].cells[0].text = 'Affected Host(s)'
            set_font(vuln_table.rows[1].cells[0].paragraphs[0].runs[0], FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            affected_hosts_list = '\n'.join([f"{idx}. {v['host']}:{v['port']}/{v['protocol']}" 
                                            for idx, v in enumerate(vuln_list, 1)])
            vuln_table.rows[1].cells[1].text = affected_hosts_list
            set_cell_font(vuln_table.rows[1].cells[1], FONT_CONFIG['name'], FONT_CONFIG['size'])
            
            # CVSS Score
            vuln_table.rows[2].cells[0].text = 'CVSS Score'
            set_font(vuln_table.rows[2].cells[0].paragraphs[0].runs[0], FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            vuln_table.rows[2].cells[1].text = vuln['cvss_base_score']
            set_cell_font(vuln_table.rows[2].cells[1], FONT_CONFIG['name'], FONT_CONFIG['size'])
            
            # CVSS Vector
            vuln_table.rows[3].cells[0].text = 'CVSS Vector'
            set_font(vuln_table.rows[3].cells[0].paragraphs[0].runs[0], FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            # Use cvss3_vector if available, otherwise fall back to cvss_vector
            cvss_vector_value = vuln['cvss3_vector'] if vuln['cvss3_vector'] != 'N/A' else vuln['cvss_vector']
            vuln_table.rows[3].cells[1].text = cvss_vector_value
            set_cell_font(vuln_table.rows[3].cells[1], FONT_CONFIG['name'], FONT_CONFIG['size'])
            
            # Observation (Synopsis)
            vuln_table.rows[4].cells[0].text = 'Observation'
            set_font(vuln_table.rows[4].cells[0].paragraphs[0].runs[0], FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            vuln_table.rows[4].cells[1].text = vuln['synopsis']
            set_cell_font(vuln_table.rows[4].cells[1], FONT_CONFIG['name'], FONT_CONFIG['size'])
            
            # Impact (Description)
            vuln_table.rows[5].cells[0].text = 'Impact'
            set_font(vuln_table.rows[5].cells[0].paragraphs[0].runs[0], FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            vuln_table.rows[5].cells[1].text = vuln['description']
            set_cell_font(vuln_table.rows[5].cells[1], FONT_CONFIG['name'], FONT_CONFIG['size'])
            
            # Area of Improvement (Solution)
            vuln_table.rows[6].cells[0].text = 'Area Of\nImprovement'
            set_font(vuln_table.rows[6].cells[0].paragraphs[0].runs[0], FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            vuln_table.rows[6].cells[1].text = vuln['solution']
            set_cell_font(vuln_table.rows[6].cells[1], FONT_CONFIG['name'], FONT_CONFIG['size'])
            
            # Screenshot / Plugin Output (merged cells)
            screenshot_cell = vuln_table.rows[7].cells[0]
            screenshot_cell.merge(vuln_table.rows[7].cells[1])
            screenshot_cell.text = ''
            
            # Add "Screenshot:" in bold
            p = screenshot_cell.paragraphs[0]
            run_screenshot = p.add_run('Screenshot/POC:\n\n')
            set_font(run_screenshot, FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            
            # Add plugin output in italic
            if vuln['plugin_output'] != 'N/A' and vuln['plugin_output'].strip():
                run_output = p.add_run(vuln['plugin_output'])
                set_font(run_output, FONT_CONFIG['name'], FONT_CONFIG['size'], italic=True)
            
            # Status
            vuln_table.rows[8].cells[0].text = 'Status'
            set_font(vuln_table.rows[8].cells[0].paragraphs[0].runs[0], FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            vuln_table.rows[8].cells[1].text = 'OPEN'
            set_font(vuln_table.rows[8].cells[1].paragraphs[0].runs[0], FONT_CONFIG['name'], FONT_CONFIG['size'], bold=True)
            
            doc.add_paragraph()  # Spacing between vulnerabilities
    
    # Save document
    doc.save(output_file)
    print(f"Report saved to: {output_file}")

def open_file(filepath):
    """Open file with default application"""
    try:
        if platform.system() == 'Windows':
            os.startfile(filepath)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.run(['open', filepath])
        else:  # Linux
            subprocess.run(['xdg-open', filepath])
        print(f"Opening: {filepath}")
    except Exception as e:
        print(f"Could not auto-open file: {e}")
        print(f"Please open manually: {filepath}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python jim_nessus_parser.py <input.nessus> [output.docx]")
        print("\nCVSS v4 Severity Thresholds:")
        print(f"  Critical: {CVSS_V4_CONFIG['critical']['min']} - {CVSS_V4_CONFIG['critical']['max']}")
        print(f"  High:     {CVSS_V4_CONFIG['high']['min']} - {CVSS_V4_CONFIG['high']['max']}")
        print(f"  Medium:   {CVSS_V4_CONFIG['medium']['min']} - {CVSS_V4_CONFIG['medium']['max']}")
        print(f"  Low:      {CVSS_V4_CONFIG['low']['min']} - {CVSS_V4_CONFIG['low']['max']}")
        print(f"  Info:     {CVSS_V4_CONFIG['info']['min']}")
        print(f"\nFont Configuration:")
        print(f"  Font: {FONT_CONFIG['name']}")
        print(f"  Size: {FONT_CONFIG['size']}pt")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'nessus_report.docx'
    
    print(f"Sir Jimbet - Parsing Nessus file: {input_file}")
    print("https://github.com/jimbet/")
    vulns_by_severity = parse_nessus(input_file)
    
    print("Converting and creating the report... Please wait...")
    create_docx(vulns_by_severity, output_file)
    
    print("Done Sir!")
    
    # Auto-open the generated file
    open_file(output_file)

if __name__ == "__main__":
    main()
    
